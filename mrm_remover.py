import base64
import json
from docx import Document
import sys
# Open the document
doc = Document(sys.argv[1])

complete = False

i = 0

skips = []
broken = []

reference_map = {}

while not complete:
    complete = True

    for element in doc.element.body.iter():
        if "tag" in element.tag and element.attrib[r'{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'].startswith("MENDELEY_CITATION_v3_") and element not in skips:
            complete = False
            i += 1
            print(i)
            citation = json.loads(base64.b64decode(element.attrib[r'{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'].removeprefix("MENDELEY_CITATION_v3_")).decode('utf-8'))["manualOverride"]["citeprocText"]
            citation = reference_map.get(citation, citation)
            reference = element.getparent().getparent()
            text = reference.getprevious()
            
            if text.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr":
                text = text.getparent()

            if text.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc":
                text.append(Document().add_paragraph(citation)._element)
            elif text.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r":
                text.text = text.text + citation
            else:
                print(f"Unable to update {json.loads(base64.b64decode(element.attrib[r'{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'].removeprefix('MENDELEY_CITATION_v3_')).decode('utf-8'))['manualOverride']} - {text.tag}")
                broken.append(citation)
                skips.append(element)
                continue

            reference.getparent().remove(reference)

doc.save(sys.argv[1].replace('.docx', '_FIXED.docx'))
